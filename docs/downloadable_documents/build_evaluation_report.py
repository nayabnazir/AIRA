from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "AIRA_Project_Report_for_Evaluation.docx"


BLUE = "0B4F86"
DARK = "102033"
MUTED = "5B6F89"
FILL = "E8EEF5"
SOFT = "F4F7FB"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=DARK):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(color)


def style_table(table, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.first_child_found_in("w:tcBorders")
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = f"w:{edge}"
                element = tc_borders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tc_borders.append(element)
                element.set(qn("w:val"), "single")
                element.set(qn("w:sz"), "4")
                element.set(qn("w:space"), "0")
                element.set(qn("w:color"), "CBD5E1")
            if header and row_idx == 0:
                set_cell_shading(cell, FILL)
        if header and row_idx == 0:
            tr_pr = row._tr.get_or_add_trPr()
            tbl_header = tr_pr.find(qn("w:tblHeader"))
            if tbl_header is None:
                tbl_header = OxmlElement("w:tblHeader")
                tr_pr.append(tbl_header)
            tbl_header.set(qn("w:val"), "true")


def add_title(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AIRA Project Technical and Viva Report")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Artificial Intelligence Requirement Analyzer")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Prepared for final-year project evaluation")
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(BLUE if level < 3 else DARK)
    return p


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(DARK)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(DARK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(DARK)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, bold=True, color=BLUE)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(cells[idx], str(value))
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    style_table(table)
    doc.add_paragraph()
    return table


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("AIRA - Artificial Intelligence Requirement Analyzer")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10

    add_title(doc)

    add_heading(doc, "1. Executive Summary")
    add_para(
        doc,
        "AIRA is a web-based Artificial Intelligence Requirement Analyzer designed to help students, analysts, and software teams create Software Requirements Specification documents, generate UML diagrams, analyze SRS ambiguity, and describe uploaded UML diagrams. The system combines a browser-based interface, Node.js backend APIs, Python AI services, MySQL storage, and optional cloud AI fallback through Gemini and OpenRouter."
    )

    add_heading(doc, "2. Problem Statement")
    add_para(
        doc,
        "Students and beginner analysts often struggle to write complete SRS documents, identify ambiguous requirements, and convert requirements into clear UML diagrams. Manual preparation is time-consuming and inconsistent, especially when users are unfamiliar with IEEE-style SRS structure or UML notation. AIRA addresses this problem by providing one workspace where users can generate SRS content, review ambiguity, create UML diagrams, describe UML images, save history, and export results."
    )

    add_heading(doc, "3. Why We Selected This Project")
    add_bullets(doc, [
        "Requirement engineering is a common pain point in academic and real-world software projects.",
        "The project combines software engineering, database design, AI support, document generation, and UML modeling in one practical system.",
        "It is useful for students preparing final-year projects because it reduces repeated manual documentation work.",
        "It demonstrates full-stack development skills: frontend, backend, database, AI services, authentication, exports, history, and premium access control.",
        "It has future expansion potential through better AI models, cloud APIs, collaboration features, and professional templates."
    ])

    add_heading(doc, "4. Project Scope and Main Modules")
    add_table(
        doc,
        ["Module", "Purpose"],
        [
            ["Authentication", "User signup, login, session handling, and password hashing."],
            ["SRS Generation", "Generates structured SRS documents from a project idea, uploaded file, or selected topic."],
            ["AI Analysis", "Checks uploaded SRS text/documents for ambiguity, clarity issues, and improvement suggestions."],
            ["UML Generation", "Generates professional UML diagrams from descriptions or SRS input, with editable diagram workspace."],
            ["UML Image Description", "Analyzes uploaded or pasted UML diagrams and creates diagram-specific descriptions."],
            ["History", "Stores generated outputs so users can review or reuse previous work."],
            ["Settings and Premium", "Provides account settings, language/theme preferences, free/premium limits, and billing flow."],
            ["Export", "Supports downloadable reports and diagrams in available formats such as PDF, Word/DOC, TXT, PNG/SVG/PlantUML where applicable."]
        ],
        widths=[2.0, 4.5],
    )

    add_heading(doc, "5. Tools, Languages, and Frameworks Used")
    add_table(
        doc,
        ["Area", "Technology Used", "Role in Project"],
        [
            ["Frontend", "HTML5, CSS3, JavaScript", "Browser pages, forms, navigation, preview screens, settings, editor interactions, and responsive UI."],
            ["Frontend framework", "No React/Vue/Angular", "The frontend is built with plain HTML, CSS, and JavaScript for simpler deployment and direct control."],
            ["Design/UI", "Custom CSS", "Theme modes, responsive layout, cards, menus, forms, sidebar, ribbons, and dark/light/soft contrast appearances."],
            ["Backend", "Node.js with Express.js", "REST APIs, authentication routes, file extraction routes, billing routes, AI route orchestration, and static frontend serving."],
            ["Database", "MySQL", "Stores users, history, subscriptions, usage, generated outputs, and other persistent records."],
            ["Database access", "mysql2", "Connects Node.js backend with MySQL database."],
            ["Authentication security", "bcryptjs", "Hashes user passwords and verifies login credentials securely."],
            ["AI services", "Python", "Runs SRS generation, ambiguity checking, UML generation, and diagram description services."],
            ["Python ML libraries", "scikit-learn, joblib, numpy, pillow, graphviz", "Model loading, structured processing, image/file support, and diagram-related utilities."],
            ["Document parsing", "pdfjs-dist and Python file utilities", "Extracts text from uploaded PDF/DOC/TXT style inputs where supported."],
            ["UML rendering", "PlantUML and Kroki/PlantUML rendering route", "Converts generated PlantUML text into professional diagram previews."],
            ["Cloud AI fallback", "Gemini API and OpenRouter API", "Optional stronger AI output for UML description and generation when API keys are configured."],
            ["Billing", "Stripe-style checkout and portal flow", "Supports premium plan upgrade, subscription tracking, and billing portal integration."],
            ["Development tools", "VS Code, Node.js, npm, Python, PowerShell, MySQL Workbench", "Project editing, dependency installation, backend running, and database management."]
        ],
        widths=[1.5, 2.2, 2.8],
    )

    add_heading(doc, "6. System Architecture")
    add_para(doc, "The project follows a layered architecture:")
    add_numbered(doc, [
        "Frontend pages collect user input and display generated outputs.",
        "Node.js Express APIs receive requests and coordinate authentication, database operations, file extraction, billing, and AI service calls.",
        "Python AI services generate SRS content, analyze ambiguity, generate UML structure, and describe UML images.",
        "MySQL stores user accounts, history, subscription state, and generated records.",
        "Optional external services such as Gemini, OpenRouter, PlantUML/Kroki, and Stripe are used when configured."
    ])

    add_heading(doc, "7. AIRA_ML Folder Usage")
    add_para(
        doc,
        "The AIRA_ML folder appears to be a research and training workspace. It contains notebooks, datasets, metrics, and training artifacts such as ambiguity and UML model experiments. In the current running application, backend and frontend files do not directly import AIRA_ML. The runtime AI services load production model artifacts from backend/aira-ai/models instead."
    )
    add_table(
        doc,
        ["Folder", "Current Role", "Recommendation"],
        [
            ["AIRA_ML", "Training/research material and project evidence.", "Keep it for documentation and future model retraining, but it is not required for normal website execution."],
            ["backend/aira-ai/models", "Runtime model files used by Python AI services.", "Keep this folder with backend deployment because the application loads these artifacts."],
            ["backend/aira-ai", "Production Python AI service layer.", "This is part of the running project and should remain connected with backend routes."]
        ],
        widths=[2.0, 2.5, 2.0],
    )

    add_heading(doc, "8. Group Work Division")
    add_para(
        doc,
        "For evaluation presentation, the project work can be divided among three members as follows. Nayab's role is presented as the core integration and implementation lead, which reflects the strongest technical contribution without presenting the project as a single-person effort."
    )
    add_table(
        doc,
        ["Member", "Assigned Responsibility", "Contribution to Present in Viva"],
        [
            ["Nayab Nazir", "Project lead, full-stack integration, backend APIs, AI workflow integration, database connection, UI consistency, testing, and deployment setup.", "Explain architecture, backend/frontend integration, AI service flow, database, API fallback, premium limits, and final debugging decisions."],
            ["Laiba Arshad", "Requirement gathering, SRS structure research, sample document review, requirement categories, and testing sample project topics.", "Explain requirement engineering concepts, SRS sections, ambiguity examples, and how sample domains were validated."],
            ["Alishba Rustam", "UML documentation support, diagram examples, UI content review, test cases, export checking, and final report preparation support.", "Explain UML diagram types, diagram description expectations, export formats, and user workflow testing."]
        ],
        widths=[1.4, 2.8, 2.3],
    )

    add_heading(doc, "9. Key Features")
    add_bullets(doc, [
        "SRS document generation from project topic, brief, or uploaded reference file.",
        "Ambiguity analysis for SRS documents with detected issues and recommended corrections.",
        "UML generation for common software diagrams such as use case, class, sequence, activity, ER, and deployment-style diagrams.",
        "UML image description through uploaded or pasted diagram images.",
        "User authentication, history, settings, theme/language controls, and export support.",
        "Free and premium access model with admin accounts having unrestricted access."
    ])

    add_heading(doc, "10. Limitations and Future Work")
    add_bullets(doc, [
        "Diagram understanding depends on OCR/image quality and API/model availability.",
        "Free local logic can handle common structures, but unusual diagrams need stronger vision AI for reliable interpretation.",
        "Future work can include stronger paid vision models, collaborative editing, instructor review mode, plagiarism checks, and IEEE template export.",
        "Deployment should secure API keys, configure Stripe webhook secrets, and use environment variables instead of hard-coded credentials."
    ])

    add_heading(doc, "11. Possible Evaluator Questions with Answers")
    qa = [
        ("What is AIRA?", "AIRA is a web-based Artificial Intelligence Requirement Analyzer that helps users generate SRS documents, analyze SRS ambiguity, generate UML diagrams, and describe uploaded UML images."),
        ("What problem does your project solve?", "It reduces the difficulty and time required to prepare requirement documents and UML diagrams, especially for students and beginner analysts."),
        ("Why did you choose this project?", "Requirement engineering is essential in every software project, and many students struggle with clear, complete, and testable requirements. AIRA gives practical support in that area."),
        ("Which language is used for frontend?", "The frontend uses HTML5, CSS3, and JavaScript."),
        ("Which framework is used for frontend?", "No major frontend framework such as React, Angular, or Vue is used. The UI is implemented with plain HTML, CSS, and JavaScript."),
        ("Which language is used for backend?", "The main backend is written in JavaScript using Node.js and Express.js."),
        ("Why did you use Python?", "Python is used for AI and ML services because it has strong libraries for model loading, text processing, image handling, and AI workflows."),
        ("Which database is used?", "MySQL is used for persistent storage, with MySQL Workbench used for database management."),
        ("How does the backend connect to MySQL?", "The Node.js backend uses the mysql2 package to connect with the MySQL database."),
        ("How are passwords secured?", "Passwords are hashed using bcryptjs before storage, so plain-text passwords are not saved."),
        ("What are the main modules?", "Authentication, SRS generation, AI ambiguity analysis, UML generation, UML image description, history, settings, premium plans, and export."),
        ("How does SRS generation work?", "The system collects a project topic, brief, or file content, sends it through backend and AI services, and produces a structured SRS document."),
        ("How does ambiguity analysis work?", "The system extracts SRS text, identifies unclear or non-testable statements, and provides recommendations to make requirements measurable and specific."),
        ("What is a good requirement?", "A good requirement is clear, complete, consistent, feasible, testable, and free from vague terms."),
        ("What is ambiguity in SRS?", "Ambiguity means a requirement can be interpreted in more than one way, such as using terms like fast, user-friendly, secure, or real-time without measurable criteria."),
        ("How does UML generation work?", "The system identifies actors, entities, processes, relationships, and diagram type, then generates a structured UML representation and renders it for preview."),
        ("What is PlantUML used for?", "PlantUML is used to convert textual UML syntax into professional diagram images."),
        ("What is Kroki used for?", "Kroki can render PlantUML diagrams through an external rendering service when available."),
        ("How does UML image description work?", "The user uploads or pastes a diagram image. The system identifies its type, extracts visible elements, detects relationships or flow, and generates a structured description."),
        ("What happens if the image is blurry?", "The system should inform the user that the image is not clearly detectable and request a clearer diagram instead of forcing an incorrect description."),
        ("Why did you add Gemini and OpenRouter?", "They provide stronger AI/vision support. Gemini is attempted first, and OpenRouter can be used as fallback if Gemini is unavailable or rate-limited."),
        ("Is AIRA_ML used in the running website?", "Not directly. AIRA_ML is a training and research folder. The running app uses backend/aira-ai and backend/aira-ai/models."),
        ("Why keep AIRA_ML then?", "It is useful as evidence of model experimentation, dataset preparation, metrics, and future retraining work."),
        ("What is the purpose of the History module?", "It stores generated outputs so users can revisit, download, or reuse previous SRS and UML results."),
        ("What is the premium plan for?", "Premium unlocks unlimited or advanced usage, multi-language support, advanced analysis, multiple export formats, premium templates, and more storage/history features."),
        ("How are admin users handled?", "Admin users are configured to bypass premium limits and access all features free of cost for management and evaluation purposes."),
        ("What export formats are supported?", "The project supports formats such as PDF, Word/DOC, TXT, and diagram formats depending on the module."),
        ("Why is PDF free and other exports premium?", "PDF is a basic common output. DOCX, PPTX, Excel, and advanced templates add more processing and professional value, so they fit premium access."),
        ("What is the role of settings?", "Settings manage account details, appearance, language, premium plan information, and logout/session controls."),
        ("What is the role of language support?", "Language support helps users from different backgrounds use the app more comfortably and generate outputs in selected languages."),
        ("What are the project limitations?", "Accuracy depends on input quality, OCR clarity, AI availability, and how well the user describes the project topic."),
        ("How can the project be improved?", "By adding stronger vision models, better dataset training, collaborative editing, cloud deployment, role-based admin dashboard, and more professional export templates."),
        ("What makes your project different?", "It combines SRS generation, ambiguity checking, UML generation, UML image description, history, settings, exports, and premium access in one integrated workspace."),
        ("Why not use only templates?", "Templates are useful, but AI-based generation can adapt content to different project topics and user inputs."),
        ("What is the purpose of using MySQL?", "MySQL stores persistent records such as users, generated history, billing state, and project outputs."),
        ("How did you test the project?", "Testing involved signup/login, database connection, file upload, SRS generation, UML generation, ambiguity analysis, description generation, exports, dark/light themes, and premium limits."),
        ("What should each member explain?", "Nayab should explain integration and technical flow; Laiba should explain requirements and SRS quality; Alishba should explain UML, documentation, and testing.")
    ]
    table = add_table(doc, ["Question", "Suggested Answer"], qa, widths=[2.2, 4.3])
    for row in table.rows[1:]:
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(BLUE)

    add_heading(doc, "12. Final Conclusion")
    add_para(
        doc,
        "AIRA is a practical final-year project because it addresses a real academic and software-engineering problem. It demonstrates full-stack development, database design, AI-assisted document processing, UML modeling, export generation, and product-style premium access. The project is suitable for evaluation because it has clear modules, visible user workflows, and strong future enhancement potential."
    )

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_doc())
